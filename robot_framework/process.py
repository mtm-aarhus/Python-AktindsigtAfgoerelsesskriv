from OpenOrchestrator.orchestrator_connection.connection import OrchestratorConnection
from OpenOrchestrator.database.queues import QueueElement
import re
import os
import pandas as pd
from office365.runtime.auth.user_credential import UserCredential
from office365.sharepoint.client_context import ClientContext
from docx.shared import Pt
from docx.shared import Inches
from docx import Document
import json
import shutil
from datetime import date
import datetime
from copy import deepcopy
from docx import Document
import os
import requests
import shutil
import uuid
import AfslutSag
import GetKmdAcessToken
from urllib.parse import quote
import math


# pylint: disable-next=unused-argument
def process(orchestrator_connection: OrchestratorConnection, queue_element: QueueElement | None = None) -> None:
    """This module contains the main process of the robot."""
    def sharepoint_client(username, password, sharepoint_site_url):
        ctx = ClientContext(sharepoint_site_url).with_credentials(UserCredential(username, password))
        web = ctx.web
        ctx.load(web)
        ctx.execute_query()
        return ctx

    def upload_to_sharepoint(client: ClientContext, folder_name: str, file_path: str, folder_url: str):
        """
        Uploads a file to a specific folder in a SharePoint document library.

        :param client: Authenticated SharePoint client context
        :param folder_name: Name of the target folder within the document library
        :param file_path: Local file path to upload
        :param folder_url: SharePoint folder URL where the file should be uploaded
        """
        try:
            # Extract file name safely
            file_name = os.path.basename(file_path)

            # Define the SharePoint document library structure
            document_library = f"{folder_url.split('/', 1)[-1]}/Delte Dokumenter/Aktindsigter"
            folder_path = f"{document_library}/{folder_name}"

            # Read file into memory (Prevents closed file issue)
            with open(file_path, "rb") as file:
                file_content = file.read()  

            # Get SharePoint folder reference
            target_folder = client.web.get_folder_by_server_relative_url(folder_url)

            # Upload file using byte content
            target_folder.upload_file(file_name, file_content)
            
            # Execute request
            client.execute_query()
            orchestrator_connection.log_info(f"✅ Successfully uploaded: {file_name} to {folder_path}")

        except Exception as e:
            orchestrator_connection.log_info(f"❌ Error uploading file: {str(e)}")

    def download_file_from_sharepoint(client, sharepoint_file_url):
        '''
        Function for downloading file from sharepoint
        '''
        file_name = sharepoint_file_url.split("/")[-1]
        download_path = os.path.join(os.getcwd(), file_name)
        with open(download_path, "wb") as local_file:
            client.web.get_file_by_server_relative_path(sharepoint_file_url).download(local_file).execute_query()
        return download_path

    def check_excel_file(file_path):
        '''
        Goes through the document list and saves the data in a dictionary.
        '''
        df = pd.read_excel(file_path)
        documents = []
        if 'Gives der aktindsigt i dokumentet? (Ja/Nej/Delvis)' in df.columns and 'Begrundelse hvis nej eller delvis' in df.columns:
            for _, row in df.iterrows():
                documents.append({
                    'title': row['Dokumenttitel'],
                    'decision': row['Gives der aktindsigt i dokumentet? (Ja/Nej/Delvis)'],
                    'reason': row['Begrundelse hvis nej eller delvis'],
                    'Akt ID': row['Akt ID'],
                    'Dok ID': row['Dok ID']
                })
        return documents

    def traverse_and_check_folders(client, folder_url, results, orchestrator_connection):
        '''
        Goes through the different folders to find the excel file (ie. the document list)
        '''
        pattern = re.compile(r"([A-Za-z]\d{4}-\d{1,10}|[A-Za-z]{3}-\d{4}-\d{6})")
        folder = client.web.get_folder_by_server_relative_url(folder_url)
        client.load(folder)
        client.execute_query()

        subfolders = folder.folders
        client.load(subfolders)
        client.execute_query()

        for subfolder in subfolders:
            subfolder_name = subfolder.properties["Name"]
            subfolder_url = f"{folder_url}/{subfolder_name}"
            if re.search(pattern, subfolder_name):
                files = subfolder.files
                client.load(files)
                client.execute_query()

                for file in files:
                    if file.properties["Name"].endswith(".xlsx"):
                        file_url = f"{subfolder_url}/{file.properties['Name']}"
                        local_file_path = download_file_from_sharepoint(client, file_url)
                        document_results = check_excel_file(local_file_path)
                        results[subfolder_name] = document_results  # Ensuring it is a list
                        os.remove(local_file_path)
                        break

            traverse_and_check_folders(client, subfolder_url, results, orchestrator_connection)

    def update_document_with_besvarelse(doc_path, case_details, DeskproTitel, AnsøgerNavn, AnsøgerEmail, Afdeling, AktindsigtsDato, Beskrivelse):

        def replace_in_paragraphs(paragraphs, replacements):
            for para in paragraphs:
                full_text = "".join(run.text for run in para.runs)
                replaced = False
                for key, val in replacements.items():
                    if key in full_text:
                        full_text = full_text.replace(key, val)
                        replaced = True

                if replaced:
                    # Ryd alle eksisterende runs
                    for run in para.runs:
                        run.text = ""
                    # Genskab som ét run med standardformat
                    para.runs[0].text = full_text


        def replace_in_tables(tables, replacements):
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_in_paragraphs(cell.paragraphs, replacements)

        replacements = {
            "[Deskprotitel]": DeskproTitel,
            "[Ansøgernavn]": AnsøgerNavn,
            "[Ansøgermail]": AnsøgerEmail,
            "[Afdeling]": Afdeling,
            "[Modtagelsesdato]": datetime.datetime.strptime(AktindsigtsDato, "%Y-%m-%dT%H:%M:%SZ").strftime("%d-%m-%Y"),
            "[beskrivelse]": Beskrivelse
        }

        doc = Document(doc_path)

        # 1. Brødtekst og tabeller i hoveddokumentet
        replace_in_paragraphs(doc.paragraphs, replacements)
        replace_in_tables(doc.tables, replacements)

        # 2. Sidehoveder og sidefødder (alle variationer)
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                replace_in_paragraphs(header.paragraphs, replacements)
                replace_in_tables(header.tables, replacements)

            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                replace_in_paragraphs(footer.paragraphs, replacements)
                replace_in_tables(footer.tables, replacements)

        doc.save("Afgørelse.docx")
        orchestrator_connection.log_info("✅ Dokument opdateret og gemt som 'Afgørelse.docx'")

    def update_internal_template_with_documenttypes(source_doc_path: str, reasons: list, placeholder: str = "[Dokumenttype]"):
        """
        Erstatter placeholder [Dokumenttype] i et dokument med en bulletliste over relevante interne dokumenttyper.
        Tilføjer visuel indrykning og anvender ikke styles, da de ikke altid er defineret.
        """
        from docx import Document
        doc = Document(source_doc_path)

        internt_reason_to_text = {
            "Internt dokument - ufærdigt arbejdsdokument": "Udkast til dokumenter",
            "Internt dokument - foreløbige og sagsforberedende overvejelser": "Dokumenter med foreløbige, interne overvejelser",
            "Internt dokument - del af intern beslutningsproces": "Dokumenter, som er indgået i en intern beslutningsproces"
        }

        relevant_texts = {
            internt_reason_to_text[r]
            for r in reasons
            if r in internt_reason_to_text
        }

        if not relevant_texts:
            orchestrator_connection.log_info("ℹ️  Ingen interne dokumenttyper fundet – ingen ændringer foretaget.")
            return

        orchestrator_connection.log_info(f"📝  Indsætter dokumenttyper i {source_doc_path}: {sorted(relevant_texts)}")

        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                parent = paragraph._element.getparent()
                insert_index = parent.index(paragraph._element)
                parent.remove(paragraph._element)

                for text in sorted(relevant_texts):
                    p = doc.add_paragraph()
                    run = p.add_run(f"• {text}")
                    run.font.size = Pt(10)
                    p.paragraph_format.left_indent = Inches(0.5)
                    parent.insert(insert_index, p._element)
                    insert_index += 1
                break

        doc.save(source_doc_path)


    def replace_placeholder_with_multiple_documents(target_doc_path: str, reason_doc_map: dict, placeholder: str):
        """
        Erstatter placeholder i target_doc med indhold fra flere dokumenter i rækkefølge.
        Hvis reason_doc_map er tom, fjernes placeholderen stille og roligt.
        """
        orchestrator_connection.log_info(f"➡️  Åbner hoveddokument for fletning: {target_doc_path}")
        target_doc = Document(target_doc_path)

        if not reason_doc_map:
            orchestrator_connection.log_info(f"➡️  Ingen dokumenter at indsætte. Fjerner placeholder '{placeholder}'")
            for paragraph in target_doc.paragraphs:
                if placeholder in paragraph.text:
                    paragraph.text = ""
            target_doc.save(target_doc_path)
            return

        for paragraph in target_doc.paragraphs:
            if placeholder in paragraph.text:
                parent = paragraph._element.getparent()
                insert_index = parent.index(paragraph._element)
                parent.remove(paragraph._element)

                for reason, doc_path in reason_doc_map.items():
                    orchestrator_connection.log_info(f"    ↪️ Indsætter '{doc_path}' pga. begrundelse: '{reason}'")

                    if not os.path.exists(doc_path):
                        print(f"    ⚠️  Fil ikke fundet: {doc_path}")
                        continue

                    source_doc = Document(doc_path)
                    for para in source_doc.paragraphs:
                        para_copy = deepcopy(para._element)
                        parent.insert(insert_index, para_copy)
                        insert_index += 1

                break
                # Ryd op i midlertidige dokumenter
        for path in set(used_doc_map.values()):
            if os.path.basename(path).startswith("temp_internal_") and os.path.exists(path):
                try:
                    os.remove(path)
                    orchestrator_connection.log_info(f"🗑  Slettede midlertidig fil: {path}")
                except Exception as e:
                    orchestrator_connection.log_info(f"⚠️  Kunne ikke slette {path}: {e}")

            target_doc.save(target_doc_path)
        orchestrator_connection.log_info(f"✅  Fletning afsluttet og gemt i: {target_doc_path}")

    def prepare_internal_document_if_needed(reasons: list, lovgivning: str, doc_map_by_lovgivning: dict) -> dict:
        """
        Finder og tilpasser internt dokument hvis nødvendigt. Returnerer mapping med den tilpassede sti.
        Funktionen undgår at lave en midlertidig kopi og arbejder direkte med originalstien.
        """
        internal_alias = "__intern__"
        updated_docs = {}

        internal_reasons = {
            "Internt dokument - ufærdigt arbejdsdokument",
            "Internt dokument - foreløbige og sagsforberedende overvejelser",
            "Internt dokument - del af intern beslutningsproces"
        }

        # Find de faktiske interne begrundelser i den oprindelige liste
        used_internal_reasons = [
            doc["reason"] for r in results.values()
            for doc in r
            if doc["decision"] in ["Nej", "Delvis"]
            and doc["reason"] in internal_reasons
        ]


        # Hvis alias er brugt, skal vi bygge og tilpasse dokumentet
        if internal_alias in reasons and used_internal_reasons:
            internal_template_key = "Internt dokument - ufærdigt arbejdsdokument"
            original_path = doc_map_by_lovgivning.get(lovgivning, {}).get(internal_template_key)

            orchestrator_connection.log_info(f"➡️  Der skal bruges internt dokument for alias: {internal_alias}")

            if original_path:
                update_internal_template_with_documenttypes(original_path, used_internal_reasons)
                updated_docs[internal_alias] = original_path
            else:
                orchestrator_connection.log_info(f"⚠️  Dokument ikke fundet: {original_path}")

        return updated_docs
    
    def extract_unique_reasons(results_dict):
        """
        Returnerer en liste med unikke begrundelser (reason) fra results,
        hvor interne begrundelser samles til én fælles type for at undgå duplikering.
        """
        internal_alias = "__intern__"
        internal_reasons = {
            "Internt dokument - ufærdigt arbejdsdokument",
            "Internt dokument - foreløbige og sagsforberedende overvejelser",
            "Internt dokument - del af intern beslutningsproces"
        }

        cleaned = set()
        for docs in results_dict.values():
            for doc in docs:
                if doc["decision"] in ["Nej", "Delvis"] and doc["reason"] and str(doc["reason"]).strip():
                    if isinstance(reason_raw, float) and math.isnan(reason_raw):
                        orchestrator_connection.log_info('Ingen begrundelse valgt')
                        reason = 'Intet valgt'
                    reason = doc["reason"].strip()
                    if reason in internal_reasons:
                        cleaned.add(internal_alias)
                    else:
                        cleaned.add(reason)

        return list(cleaned)

    #Getting queue elements
    queue_json = json.loads(queue_element.data)
    DeskproTitel = queue_json.get('Aktindsigtsovermappe') or ""
    AnsøgerNavn = queue_json.get('AnsøgerNavn') or ""
    AnsøgerEmail = queue_json.get('AnsøgerEmail') or ""
    Afdeling = queue_json.get('Afdeling') or ""
    DeskProID = queue_json.get('DeskProID')
    KMDNovaURL = orchestrator_connection.get_constant("KMDNovaURL").value
    AktindsigtsDato = queue_json.get("AktindsigtsDato") or ""
    Lovgivning = queue_json.get('Lovgivning') or None

    #Getting oo stuff
    API_aktbob = orchestrator_connection.get_credential('AktbobAPIKey')
    url = f'{API_aktbob.username}/cases?deskproId={DeskProID}'
    key = API_aktbob.password

    #Getting description of aktindsigt
    headers = {
        'ApiKey': key
        }
    response = requests.request("GET", url, headers=headers)
    data = response.json()
    Beskrivelse = data[0].get("requestDescription", "")
    if not Beskrivelse:
        Beskrivelse = ""

    RobotCredentials = orchestrator_connection.get_credential("RobotCredentials")
    username = RobotCredentials.username
    password = RobotCredentials.password
    sharepoint_site_url = orchestrator_connection.get_constant("AktbobSharePointURL").value
    parent_folder_url = sharepoint_site_url.split(".com")[-1] +'/Delte Dokumenter/'

    if Afdeling != 'Plan og Byggeri':
        if Lovgivning == "Ikke part, miljøoplysning (1985 offentligthedsloven og miljøoplysningsloven)":
            doc_path = r'AB-hovedfrase - Helt eller delvist afslag - OFL og MOL.docx'
        elif Lovgivning == "Part, miljøoplysning (2012 forvaltningsloven og miljøoplysningsloven)":
            doc_path = r'AB-hovedfrase - Helt eller delvist afslag - FVL og MOL.docx'
        elif Lovgivning == "Part, ingen miljøoplysning (2014 forvaltningsloven)":
            doc_path = r'AB-hovedfrase - Helt eller delvist afslag - FVL - ikke MOL.docx'
        elif Lovgivning == "Ikke part, ingen miljøoplysning (2020 offentlighedsloven)":
            doc_path = r'AB-hovedfrase - helt eller delvist afslag - OFL - ikke MOL.docx'
        elif Lovgivning == "Andet (Genererer fuld frase)":
            doc_path = r'MISSING.docx'
        else: 
            doc_path = r'MISSING.docx'
    else:
        if Lovgivning == "Ikke part, miljøoplysning (1985 offentligthedsloven og miljøoplysningsloven)":
            doc_path = r'AB-hovedfrase - Helt eller delvist afslag - OFL og MOL.docx'
        elif Lovgivning == "Part, miljøoplysning (2012 forvaltningsloven og miljøoplysningsloven)":
            doc_path = r'AB-hovedfrase - Helt eller delvist afslag - FVL og MOL.docx'
        elif Lovgivning == "Part, ingen miljøoplysning (2014 forvaltningsloven)":
            doc_path = r'AB-hovedfrase - Helt eller delvist afslag - FVL - ikke MOL.docx'
        elif Lovgivning == "Ikke part, ingen miljøoplysning (2020 offentlighedsloven)":
            doc_path = r'AB-hovedfrase - helt eller delvist afslag - OFL - ikke MOL.docx'
        elif Lovgivning == "Andet (Genererer fuld frase)":
            doc_path = r'MISSING.docx'
        else: 
            doc_path = r'MISSING.docx'

    doc_map_by_lovgivning = {
        "Ikke part, miljøoplysning (1985 offentligthedsloven og miljøoplysningsloven)": {
            "Internt dokument - ufærdigt arbejdsdokument": "AB-minifrase - internt dokument - OFL og MOL.docx",
            "Internt dokument - foreløbige og sagsforberedende overvejelser": "AB-minifrase - internt dokument - OFL og MOL.docx",
            "Internt dokument - del af intern beslutningsproces": "AB-minifrase - internt dokument - OFL og MOL.docx",
            "Særlige dokumenter - korrespondance med sagkyndig rådgiver vedr. tvistsag": "AB-minifrase - sagkyndig rådgivning - OFL og MOL.docx",
            "Særlige dokumenter - statistik og undersøgelser": "MISSING.docx",
            "Særlige dokumenter - straffesag": "AB-minifrase - Dokument i straffesag - OFL og MOL.docx",
            "Tavshedsbelagte oplysninger - om private forhold": "AB-minifrase - Private forhold - OFL og MOL.docx",
            "Tavshedsbelagte oplysninger - forretningsforhold": "AB-minifrase - Forretningsforhold - OFL og MOL.docx",
            "Tavshedsbelagte oplysninger - Andet (uddybes i afgørelsen)": "MISSING.docx",
            " ": "MISSING.docx",
            "Intet valgt": "Ingen begrundelse valgt.docx"
        },
        "Part, miljøoplysning (2012 forvaltningsloven og miljøoplysningsloven)": {
            "Internt dokument - ufærdigt arbejdsdokument": "AB-minifrase - internt dokument - FVL og MOL.docx",
            "Internt dokument - foreløbige og sagsforberedende overvejelser": "AB-minifrase - internt dokument - FVL og MOL.docx",
            "Internt dokument - del af intern beslutningsproces": "AB-minifrase - internt dokument - FVL og MOL.docx",
            "Særlige dokumenter - korrespondance med sagkyndig rådgiver vedr. tvistsag": "AB-minifrase - sagkyndig rådgivning - FVL og MOL.docx",
            "Særlige dokumenter - statistik og undersøgelser": "MISSING.docx",
            "Særlige dokumenter - straffesag": "AB-minifrase - Dokument i straffesag - FVL og MOL.docx",
            "Tavshedsbelagte oplysninger - om private forhold": "AB-minifrase - Private forhold - FVL og MOL.docx",
            "Tavshedsbelagte oplysninger - forretningsforhold": "AB-minifrase - Forretningsforhold - FVL og MOL.docx",
            "Tavshedsbelagte oplysninger - Andet (uddybes i afgørelsen)": "MISSING.docx",
            " ": "MISSING.docx",
            "Intet valgt": "Ingen begrundelse valgt.docx"
        },
        "Part, ingen miljøoplysning (2014 forvaltningsloven)": {
            "Internt dokument - ufærdigt arbejdsdokument": "AB-minifrase - internt dokument - FVL.docx",
            "Internt dokument - foreløbige og sagsforberedende overvejelser": "AB-minifrase - internt dokument - FVL.docx",
            "Internt dokument - del af intern beslutningsproces": "AB-minifrase - internt dokument - FVL.docx",
            "Særlige dokumenter - korrespondance med sagkyndig rådgiver vedr. tvistsag": "AB-minifrase - Sagkyndig rådgivning - FVL.docx",
            "Særlige dokumenter - statistik og undersøgelser": "MISSING.docx",
            "Særlige dokumenter - straffesag": "AB-minifrase - Dokument i straffesag - FVL.docx",
            "Tavshedsbelagte oplysninger - om private forhold": "AB-minifrase - Private forhold - FVL.docx",
            "Tavshedsbelagte oplysninger - forretningsforhold": "AB-minifrase - Forretningsforhold - FVL.docx",
            "Tavshedsbelagte oplysninger - Andet (uddybes i afgørelsen)": "MISSING.docx",
            " ": "MISSING.docx",
            "Intet valgt": "Ingen begrundelse valgt.docx"
        },
        "Ikke part, ingen miljøoplysning (2020 offentlighedsloven)": {
            "Internt dokument - ufærdigt arbejdsdokument": "AB-minifrase - Internt dokument - OFL.docx",
            "Internt dokument - foreløbige og sagsforberedende overvejelser": "AB-minifrase - Internt dokument - OFL.docx",
            "Internt dokument - del af intern beslutningsproces": "AB-minifrase - Internt dokument - OFL.docx",
            "Særlige dokumenter - korrespondance med sagkyndig rådgiver vedr. tvistsag": "AB-minifrase - Sagkyndig rådgivning - OFL.docx",
            "Særlige dokumenter - statistik og undersøgelser": "MISSING.docx",
            "Særlige dokumenter - straffesag": "AB-minifrase - Dokument i straffesag - OFL.docx",
            "Tavshedsbelagte oplysninger - om private forhold": "AB-minifrase - Private forhold - OFL.docx",
            "Tavshedsbelagte oplysninger - forretningsforhold": "AB-minifrase - Forretningsforhold - OFL.docx",
            "Tavshedsbelagte oplysninger - Andet (uddybes i afgørelsen)": "MISSING.docx",
            " ": "MISSING.docx",
            "Intet valgt": "Ingen begrundelse valgt.docx"
        },
        "Ved ikke (Genererer fuld frase)": {
            "Internt dokument - ufærdigt arbejdsdokument": "AB-minifrase - internt dokument - OFL og MOL.docx",
            "Internt dokument - foreløbige og sagsforberedende overvejelser": "AB-minifrase - internt dokument - OFL og MOL.docx",
            "Internt dokument - del af intern beslutningsproces": "AB-minifrase - internt dokument - OFL og MOL.docx",
            "Særlige dokumenter - korrespondance med sagkyndig rådgiver vedr. tvistsag": "MISSING.docx",
            "Særlige dokumenter - statistik og undersøgelser": "MISSING.docx",
            "Særlige dokumenter - straffesag": "MISSING.docx",
            "Tavshedsbelagte oplysninger - om private forhold": "AB-minifrase - Private forhold - OFL og MOL.docx",
            "Tavshedsbelagte oplysninger - forretningsforhold": "AB-minifrase - Forretningsforhold - OFL og MOL.docx",
            "Tavshedsbelagte oplysninger - Andet (uddybes i afgørelsen)": "MISSING.docx",
            " ": "MISSING.docx",
            "Intet valgt": "Ingen begrundelse valgt.docx"
        }
    }

    #Skal den lokale version af afgørelse slettes?
    slet = True

    client = sharepoint_client(username, password, sharepoint_site_url)
    results = {}
    traverse_and_check_folders(client, f'{parent_folder_url}Dokumentlister/{DeskproTitel}', results, orchestrator_connection)
    update_document_with_besvarelse(doc_path, results, DeskproTitel= DeskproTitel, AnsøgerEmail= AnsøgerEmail, AnsøgerNavn= AnsøgerNavn, Afdeling= Afdeling, AktindsigtsDato = AktindsigtsDato, Beskrivelse = Beskrivelse)

    unique_reasons = extract_unique_reasons(results)
    internal_docs = prepare_internal_document_if_needed(unique_reasons, Lovgivning, doc_map_by_lovgivning)
    used_doc_map = {}

    for reason in unique_reasons:
        if reason == "__intern__" and "__intern__" in internal_docs:
            used_doc_map[reason] = internal_docs["__intern__"]
        else:
            doc = doc_map_by_lovgivning.get(Lovgivning, {}).get(reason)
            if doc:
                used_doc_map[reason] = doc

    replace_placeholder_with_multiple_documents("Afgørelse.docx", used_doc_map, "[RELEVANTE_TEKSTER]")

    orchestrator_connection.log_info('Setting cases as finished in nova if novacase')
    KMD_access_token = GetKmdAcessToken.GetKMDToken(orchestrator_connection = orchestrator_connection)
    AfslutSag.invoke_AfslutSag(KMDNovaURL, KMD_access_token, DeskProID= DeskProID, orchestrator_connection= orchestrator_connection)

    orchestrator_connection.log_info('Document updating, uploading to sharepoint')
    upload_to_sharepoint(client, DeskproTitel, r'Afgørelse.docx', folder_url = f'{parent_folder_url}Aktindsigter/{DeskproTitel}')
    if slet:
        afgorelse_path = "Afgørelse.docx"
        if os.path.exists(afgorelse_path):
            os.remove(afgorelse_path)
            orchestrator_connection.log_info(f"🗑  Slettede midlertidig fil: {afgorelse_path}")
        else:
            orchestrator_connection.log_info(f"⚠️  Filen '{afgorelse_path}' blev ikke fundet og kunne derfor ikke slettes.")
    #Putting sharepointlink to case top folder in deskpro

    deskproURL = orchestrator_connection.get_constant('DeskproOvermappeAPILink').value
    API_url = orchestrator_connection.get_constant("AktbobSharePointURL").value

    payload = json.dumps({
        "deskproTicketId": f'{DeskProID}',
        "overmappeURL": f'{API_url}/Delte%20Dokumenter/Aktindsigter/{quote(DeskproTitel)}'
        })
    
    headers = {
        'Content-Type': 'application/json'
        }
    response_deskpro = requests.request("POST", deskproURL, headers=headers, data=payload)
    response_deskpro.raise_for_status()